import io
import fitz  # PyMuPDF


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extracts text from a PDF file provided as bytes (PyMuPDF)."""
    text_content = []
    try:
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            clean_text = " ".join(text.split())
            if clean_text:
                text_content.append(f"--- Page {page_num + 1} ---\n{clean_text}\n")
        return "\n".join(text_content)
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        raise e


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extracts text (paragraphs + table cells) from a .docx file provided as bytes."""
    from docx import Document  # python-docx

    try:
        document = Document(io.BytesIO(docx_bytes))
        blocks = []

        for para in document.paragraphs:
            t = para.text.strip()
            if t:
                blocks.append(t)

        # Capture tabular content so technical docs aren't silently dropped
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        return "\n".join(blocks)
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        raise e


def extract_text_from_document(file_bytes: bytes, filename: str) -> str:
    """
    Dispatches to the correct extractor based on the file extension.
    Previously every upload was forced through the PDF parser (filetype="pdf"),
    which silently broke .docx ingestion.
    """
    name = (filename or "").lower()
    if name.endswith(".docx"):
        return extract_text_from_docx_bytes(file_bytes)
    if name.endswith(".pdf"):
        return extract_text_from_pdf_bytes(file_bytes)
    # Unknown extension: best-effort PDF parse
    return extract_text_from_pdf_bytes(file_bytes)
